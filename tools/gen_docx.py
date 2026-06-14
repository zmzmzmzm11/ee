"""Generate EveBox deployment document as Word .docx (local project version)"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ---- Style setup ----
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

def add_heading_styled(text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '微软雅黑'
    return h

def add_para(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    if bold:
        run.bold = True
    return p

def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    for line in text.split('\n'):
        if line.strip():
            run = p.add_run(line + '\n')
        else:
            run = p.add_run('\n')
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run.font.size = Pt(10)
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.5 * (level + 1))
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.name = '微软雅黑'
        run.font.size = Pt(9)
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r+1].cells[c]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            run.font.name = '微软雅黑'
            run.font.size = Pt(9)
    doc.add_paragraph()
    return table

# ==================== COVER ====================
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('EveBox')
run.bold = True
run.font.size = Pt(36)
run.font.name = '微软雅黑'
run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Suricata EVE 事件管理与告警系统')
run.font.size = Pt(16)
run.font.name = '微软雅黑'
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run('安装部署与使用文档')
run.font.size = Pt(14)
run.font.name = '微软雅黑'

doc.add_paragraph()
ver = doc.add_paragraph()
ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = ver.add_run('版本 0.25.0-dev | 2026年6月')
run.font.size = Pt(11)
run.font.name = '微软雅黑'
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_paragraph()
repo = doc.add_paragraph()
repo.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = repo.add_run('GitHub: https://github.com/zmzmzmzm11/ee')
run.font.size = Pt(10)
run.font.name = 'Consolas'
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ==================== TOC ====================
add_heading_styled('目录', 1)
toc_items = [
    '1. 项目简介', '2. 本地开发环境', '3. 快速启动', '4. 从源码编译',
    '5. 配置说明', '6. 运行模式', '7. 用户认证与角色管理',
    '8. 管理员功能', '9. CLI 命令参考', '10. 项目目录结构', '11. 常见问题'
]
for item in toc_items:
    add_para(item)
doc.add_page_break()

# ==================== 1. 项目简介 ====================
add_heading_styled('1. 项目简介', 1)

add_para('EveBox 是一个基于 Web 的 Suricata EVE 事件查看器和告警管理系统，使用 Rust 编写后端，SolidJS 编写前端。本项目基于 Jason Ish 的 evebox 进行二次开发，增加了用户角色管理、管理员功能等特性。')

add_table(
    ['属性', '值'],
    [
        ['GitHub 仓库', 'https://github.com/zmzmzmzm11/ee'],
        ['上游仓库', 'https://github.com/jasonish/evebox'],
        ['许可证', 'MIT'],
        ['当前版本', '0.25.0-dev'],
        ['后端语言', 'Rust (edition 2024, 最低 1.85.0)'],
        ['前端框架', 'SolidJS + Vite + Bootstrap 5'],
        ['Web 框架', 'Axum 0.8'],
        ['数据库支持', 'Elasticsearch 7.10+ / 嵌入式 SQLite'],
    ]
)

add_para('核心功能：', bold=True)
add_bullet('Web 事件查看器，采用"收件箱"（Inbox）模式管理告警')
add_bullet('事件搜索、过滤、归档、升级（Escalate）')
add_bullet('嵌入式 SQLite 支持，无需外部数据库即可运行')
add_bullet('用户认证与基于角色的权限管理（admin / user）')
add_bullet('管理员控制台：用户管理、自动归档过滤器、ES 索引管理')
add_bullet('JA4 TLS 指纹识别')
add_bullet('GeoIP 地理信息支持')
add_bullet('多种部署模式：Server、Agent、One-shot、Docker')

doc.add_page_break()

# ==================== 2. 本地开发环境 ====================
add_heading_styled('2. 本地开发环境', 1)

add_heading_styled('2.1 环境概览', 2)
add_para('本项目在 Windows 10 上开发，使用以下组件构建完整的开发测试环境：')

add_table(
    ['组件', '版本/路径', '端口', '说明'],
    [
        ['Elasticsearch', '7.17.28 (本地目录 elasticsearch/)', '9200', '事件数据存储'],
        ['Suricata 模拟器', 'tools/suricata_simulator.py', '-', '生成测试 EVE 事件'],
        ['EveBox Server', 'target/debug/evebox.exe', '5636 (自动选择)', 'Web 管理界面'],
        ['Java JDK', 'C:\\Program Files\\Java\\jdk-17.0.18', '-', '运行 ES 所需'],
        ['Python', '系统安装', '-', '运行模拟器'],
        ['Node.js', 'v18+', '3636 (dev)', '前端构建与开发'],
        ['Rust', '1.85.0+', '-', '后端编译'],
        ['Git Bash', 'MSYS2', '-', 'Shell 环境'],
    ]
)

add_heading_styled('2.2 必备依赖安装', 2)

add_para('1) 安装 Rust：', bold=True)
add_code('curl --proto "=https" --tlsv1.2 -sSf https://sh.rustup.rs | sh\nrustup default stable\nrustc --version   # 应 >= 1.85.0')

add_para('2) 安装 Node.js：', bold=True)
add_para('从 https://nodejs.org 下载 v18+ LTS 版本并安装。')
add_code('node --version    # 应 >= v18\nnpm --version')

add_para('3) 安装 Java JDK 17（仅运行 Elasticsearch 时需要）：', bold=True)
add_para('从 https://adoptium.net 下载 JDK 17，安装到 C:\\Program Files\\Java\\jdk-17.0.18。')

add_para('4) 安装 Python 3（仅使用模拟器时需要）：', bold=True)
add_para('从 https://python.org 下载并安装，确保 python 命令可用。')

add_heading_styled('2.3 克隆项目', 2)
add_code('git clone git@github.com:zmzmzmzm11/ee.git\ncd ee')

add_para('如果 GitHub SSH 不通，可以通过 SSH 443 端口：', bold=False)
add_code('''# 在 ~/.ssh/config 中添加：
Host github.com
    HostName ssh.github.com
    Port 443
    User git
    IdentityFile ~/.ssh/id_ed25519''')

doc.add_page_break()

# ==================== 3. 快速启动 ====================
add_heading_styled('3. 快速启动', 1)

add_heading_styled('3.1 一键启动（推荐）', 2)
add_para('项目提供 PowerShell 脚本一键启动全部组件。在项目根目录打开 PowerShell：')
add_code('.\\control.ps1 start')

add_para('脚本会自动执行：', bold=False)
add_bullet('检查 Java、evebox.exe、模拟器脚本是否就绪')
add_bullet('启动 Elasticsearch（如未运行）并等待就绪')
add_bullet('启动 Suricata EVE 事件模拟器（生成测试数据）')
add_bullet('启动 EveBox Server 并自动打开浏览器')

add_para('其他控制命令：')
add_code('.\\control.ps1 stop       # 停止所有组件\n.\\control.ps1 restart    # 重启所有组件\n.\\control.ps1 status     # 查看运行状态\n.\\control.ps1            # 交互式菜单')

add_heading_styled('3.2 Git Bash 启动', 2)
add_code('./control.sh start    # 一键启动\n./control.sh stop     # 一键停止\n./control.sh status   # 查看状态')

add_heading_styled('3.3 独立启动脚本', 2)
add_para('也可以使用独立的启动/停止脚本：')
add_code('.\\start-all.ps1    # 启动全部组件（PowerShell）\n.\\stop-all.ps1     # 停止全部组件（PowerShell）')

add_heading_styled('3.4 启动后的验证', 2)
add_para('启动完成后可以用以下命令验证各组件是否正常：')
add_code('''# 检查 ES 索引
curl http://localhost:9200/_cat/indices/evebox*?format=json

# 检查 EveBox API
curl http://127.0.0.1:5636/api/version

# 查看告警数据
curl http://127.0.0.1:5636/api/alerts | python -m json.tool

# 查看传感器列表
curl http://127.0.0.1:5636/api/sensors

# 查看事件类型
curl http://127.0.0.1:5636/api/event_types''')

add_para('浏览器访问 http://127.0.0.1:5636 即可打开 Web 管理界面。')

doc.add_page_break()

# ==================== 4. 从源码编译 ====================
add_heading_styled('4. 从源码编译', 1)

add_heading_styled('4.1 构建前端', 2)
add_code('''cd webapp
npm ci                  # 安装依赖
npm run build           # 构建生产版本到 dist/''')

add_para('前端资源（编译后的 JS/CSS/HTML）会被复制到 resources/webapp/，Rust 编译时通过 rust-embed 嵌入二进制。')

add_heading_styled('4.2 构建后端', 2)
add_code('''# 生成 git 版本信息
cd webapp
echo "export const GIT_REV = \\"$(git rev-parse --short HEAD)\\";" > src/gitrev.ts
npm run build
cd ..

# 复制前端产物
rm -rf resources/webapp
cp -a webapp/dist resources/webapp

# 编译 Rust
cargo build              # Debug 版本，快速编译
cargo build --release    # Release 版本，优化性能''')

add_para('二进制输出位置：')
add_bullet('Debug: target/debug/evebox.exe')
add_bullet('Release: target/release/evebox.exe')

add_heading_styled('4.3 使用 Makefile（推荐）', 2)
add_code('make          # 构建前端 + 编译 debug 版本\nmake dist     # 构建前端 + 编译 release 版本 + 打包 zip')

add_heading_styled('4.4 前端开发模式', 2)
add_para('前端支持热更新开发模式（修改代码后浏览器自动刷新）：')
add_code('cd webapp\nnpm run dev    # 启动 Vite 开发服务器，端口 3636')

add_para('Vite 开发服务器会将 /api 请求代理到 http://127.0.0.1:5636（后端服务器）。开发时需同时运行后端：')
add_code('# 终端 1: 启动后端\n./target/debug/evebox server --datastore elasticsearch --no-auth --no-tls -D ./data -e http://localhost:9200 -i evebox --input ./data/suricata-eve.json\n\n# 终端 2: 启动前端开发服务器\ncd webapp && npm run dev\n\n# 浏览器访问 http://localhost:3636')

doc.add_page_break()

# ==================== 5. 配置说明 ====================
add_heading_styled('5. 配置说明', 1)

add_para('EveBox 支持三种配置方式（优先级从高到低）：')
add_bullet('命令行参数：直接在命令行中指定')
add_bullet('配置文件：使用 YAML 格式的配置文件（-c 参数指定）')
add_bullet('环境变量：通过环境变量设置')

add_heading_styled('5.1 本项目的典型启动命令', 2)
add_para('这是本项目开发环境中使用的启动命令（见 control.ps1）：')
add_code('''./target/debug/evebox.exe server \\\n  --datastore elasticsearch \\\n  --no-auth --no-tls \\\n  -D ./data \\\n  -e http://localhost:9200 \\\n  -i evebox \\\n  --input ./data/suricata-eve.json \\\n  -p 5636''')

add_para('各参数说明：')
add_table(
    ['参数', '值', '说明'],
    [
        ['--datastore', 'elasticsearch', '使用 Elasticsearch 作为事件存储'],
        ['--no-auth', '-', '禁用用户认证（开发模式）'],
        ['--no-tls', '-', '禁用 TLS（本地开发）'],
        ['-D', './data', '数据目录（存储 ES 索引、bookmark 等）'],
        ['-e', 'http://localhost:9200', 'Elasticsearch 地址'],
        ['-i', 'evebox', 'ES 索引前缀（索引名: evebox-YYYY.MM.DD）'],
        ['--input', './data/suricata-eve.json', 'Suricata EVE 输入文件'],
        ['-p', '5636', 'Web 服务端口'],
    ]
)

add_heading_styled('5.2 配置文件方式', 2)
add_para('创建 evebox.yaml 配置文件后使用 -c 参数加载：')
add_code('''data-directory: ./data

http:
  host: 127.0.0.1
  port: 5636

authentication:
  required: true         # 开启用户认证

database:
  type: elasticsearch
  elasticsearch:
    url: http://localhost:9200
    index: evebox        # 索引前缀

input:
  enabled: true
  paths:
    - ./data/suricata-eve.json''')

add_para('使用配置文件启动：')
add_code('./target/debug/evebox server -c evebox.yaml')

add_heading_styled('5.3 主要环境变量', 2)
add_table(
    ['环境变量', '说明', '默认值'],
    [
        ['EVEBOX_HTTP_HOST', '绑定的主机名/IP', '127.0.0.1'],
        ['EVEBOX_HTTP_PORT', '绑定的端口', '5636'],
        ['EVEBOX_DATA_DIRECTORY', '数据目录', '-'],
        ['EVEBOX_ELASTICSEARCH_URL', 'ES URL', 'http://localhost:9200'],
        ['EVEBOX_ELASTICSEARCH_INDEX', 'ES 索引前缀', 'logstash'],
        ['EVEBOX_AUTHENTICATION_REQUIRED', '开启认证', '-'],
        ['EVE_OUTPUT', '模拟器输出路径', 'data/suricata-eve.json'],
        ['EVE_INTERVAL', '模拟器事件间隔(秒)', '1.5'],
    ]
)

doc.add_page_break()

# ==================== 6. 运行模式 ====================
add_heading_styled('6. 运行模式', 1)

add_heading_styled('6.1 开发模式：ES + Suricata 模拟器 + EveBox', 2)
add_para('本项目默认的开发运行模式，三组件联动：')
add_code('''┌─────────────────┐     ┌──────────────────┐     ┌──────────────┐
│ Suricata 模拟器  │────▶│  Elasticsearch   │◀────│   EveBox     │
│ (EVE JSON 生成)  │     │  (localhost:9200) │     │ (Web :5636)  │
└─────────────────┘     └──────────────────┘     └──────────────┘''')

add_bullet('Suricata 模拟器持续生成 EVE 事件到 data/suricata-eve.json（约 1 条/1.5 秒）')
add_bullet('EveBox 读取 EVE 文件并写入 Elasticsearch（索引前缀 evebox）')
add_bullet('EveBox 同时提供 Web 界面读取 ES 中的事件数据')

add_heading_styled('6.2 生产模式：Server + Elasticsearch', 2)
add_para('适用于已有 Suricata + Elasticsearch + Filebeat/Logstash 的生产环境：')
add_code('./evebox server -e http://elasticsearch:9200 -i logstash')

add_heading_styled('6.3 独立模式：Server + 嵌入式 SQLite', 2)
add_para('适用于中小规模部署，无需 Elasticsearch：')
add_code('./evebox server --datastore sqlite -D ./data --input /var/log/suricata/eve.json')

add_heading_styled('6.4 Agent 模式', 2)
add_para('Agent 读取 EVE 文件并转发到远程 EveBox Server 或直接写入 ES：')
add_code('./evebox agent --server http://evebox-server:5636 /var/log/suricata/eve.json')

add_heading_styled('6.5 临时分析：One-shot 模式', 2)
add_para('一次性导入 EVE JSON 文件进行临时分析：')
add_code('./evebox oneshot /path/to/eve.json')

doc.add_page_break()

# ==================== 7. 用户认证与角色管理 ====================
add_heading_styled('7. 用户认证与角色管理', 1)

add_para('本项目在原版 evebox 基础上增加了完整的用户角色管理功能，支持 admin/user 两种角色。')

add_heading_styled('7.1 首次启动自动创建管理员', 2)
add_para('当 authentication.required 为 true 且数据库中没有用户时，服务器自动创建管理员账号，控制台打印：')
add_code('Created administrator username and password: username=admin, password=<随机12位密码>')
add_para('⚠ 此密码仅在首次启动时打印一次，请务必记录！', bold=True)

add_heading_styled('7.2 命令行管理用户', 2)
add_para('通过 CLI 命令管理用户（需指定数据目录以找到 config.sqlite）：')
add_code('''# 列出所有用户
./target/debug/evebox --data-directory ./data config users ls

# 交互式添加用户
./target/debug/evebox --data-directory ./data config users add

# 指定用户名和密码添加
./target/debug/evebox --data-directory ./data config users add --username admin --password mypass

# 删除用户
./target/debug/evebox --data-directory ./data config users rm <用户名>

# 修改用户密码
./target/debug/evebox --data-directory ./data config users passwd <用户名>''')

add_heading_styled('7.3 Web 界面管理用户', 2)
add_para('以管理员账号登录后，导航栏出现 Admin 菜单，进入 Admin → Users 页面：')

add_table(
    ['操作', '步骤'],
    [
        ['创建用户', '点击 "Create User" → 填写用户名、密码（≥4字符）、选择角色 → 点击 Create'],
        ['删除用户', '在用户列表中点击对应行的 "Delete" 按钮 → 确认删除'],
        ['重置密码', '点击 "Reset Password" → 输入新密码 → 点击 Reset Password'],
    ]
)

add_para('创建用户时的验证规则：', bold=True)
add_bullet('用户名不能为空')
add_bullet('密码至少 4 个字符')
add_bullet('角色必须选择 admin 或 user')
add_bullet('表单验证失败时，错误信息会显示在弹窗内部')

add_heading_styled('7.4 角色权限详解', 2)
add_table(
    ['功能', 'admin', 'user'],
    [
        ['查看告警/事件', '✓', '✓'],
        ['搜索与过滤', '✓', '✓'],
        ['归档/升级事件', '✓', '✓'],
        ['仪表盘', '✓', '✓'],
        ['用户管理（CRUD）', '✓', '✗'],
        ['自动归档过滤器管理', '✓', '✗'],
        ['ES 索引管理', '✓', '✗'],
        ['JA4 数据库更新', '✓', '✗'],
    ]
)

add_para('权限检查在服务端进行，前端显示错误提示。若以普通用户身份执行管理员操作，将返回 "admin role required" 错误。')

doc.add_page_break()

# ==================== 8. 管理员功能 ====================
add_heading_styled('8. 管理员功能', 1)

add_para('以管理员身份登录后，导航栏出现 Admin 下拉菜单，包含以下子页面：')

add_heading_styled('8.1 Users — 用户管理', 2)
add_bullet('用户列表：显示用户名、角色标签（admin=红色, user=蓝色）')
add_bullet('创建用户：弹窗表单，支持用户名、密码、角色')
add_bullet('删除用户：点击删除按钮，弹出确认对话框')
add_bullet('重置密码：点击重置按钮，在弹窗中输入新密码')

add_heading_styled('8.2 Filters — 自动归档过滤器', 2)
add_para('配置规则使匹配的事件自动归档，减少告警噪音。')
add_bullet('过滤器条件：Sensor（传感器）、Source IP、Destination IP、Signature ID')
add_bullet('任意字段可用 * 通配符匹配所有')
add_bullet('过滤器存储在 config.sqlite 的 filters 表中')

add_heading_styled('8.3 Elasticsearch — ES 管理', 2)
add_bullet('查看 ES 集群信息和索引状态')
add_bullet('删除指定日期之前的旧索引')
add_bullet('更新 JA4 指纹数据库（TLS 指纹识别）')

# ==================== 9. CLI 命令参考 ====================
add_heading_styled('9. CLI 命令参考', 1)

add_heading_styled('9.1 server 子命令（启动服务器）', 2)
add_table(
    ['参数', '短参数', '默认值', '说明'],
    [
        ['--config <FILE>', '-c', '-', 'YAML 配置文件路径'],
        ['--host <HOST>', '-', '127.0.0.1', '绑定地址'],
        ['--port <PORT>', '-p', '5636', '绑定端口'],
        ['--datastore <TYPE>', '-', 'elasticsearch', '数据库: elasticsearch/sqlite'],
        ['--elasticsearch <URL>', '-e', 'http://localhost:9200', 'ES URL'],
        ['--index <NAME>', '-i', 'logstash', 'ES 索引前缀'],
        ['--input <FILE>', '-', '-', 'Suricata EVE 输入文件'],
        ['--data-directory <DIR>', '-D', '-', '数据目录'],
        ['--no-auth', '-', '-', '禁用认证（开发模式）'],
        ['--no-tls', '-', '-', '禁用 TLS'],
        ['--sqlite', '-', '-', '使用 SQLite 快捷方式'],
        ['--verbose', '-v', '-', '详细日志（-v=DEBUG, -vv=TRACE）'],
    ]
)

add_heading_styled('9.2 agent 子命令（事件采集代理）', 2)
add_code('''# 发送 EVE 事件到 EveBox Server
./evebox agent --server http://evebox:5636 /var/log/suricata/eve.json

# 直接写入 Elasticsearch
./evebox agent --elasticsearch -e http://es:9200 /var/log/suricata/eve.json''')

add_heading_styled('9.3 config users 子命令（用户管理）', 2)
add_code('''./evebox --data-directory ./data config users add              # 添加用户
./evebox --data-directory ./data config users ls               # 列出用户
./evebox --data-directory ./data config users rm <用户名>       # 删除用户
./evebox --data-directory ./data config users passwd <用户名>   # 修改密码''')

add_heading_styled('9.4 sqlite 常用命令', 2)
add_code('''./evebox sqlite info events.sqlite                    # 查看数据库信息
./evebox sqlite dump events.sqlite                    # 导出事件
./evebox sqlite load --input eve.json events.sqlite   # 导入 EVE 文件
./evebox sqlite fts enable events.sqlite              # 启用全文搜索
./evebox sqlite optimize events.sqlite                # 优化数据库''')

add_heading_styled('9.5 oneshot 子命令（临时分析）', 2)
add_code('./evebox oneshot eve.json          # 导入并打开浏览器分析\n./evebox oneshot --no-open eve.json  # 仅导入，不打开浏览器')

doc.add_page_break()

# ==================== 10. 项目目录结构 ====================
add_heading_styled('10. 项目目录结构', 1)

add_code('''evebox/                          # 项目根目录
├── src/                           # Rust 后端源码
│   ├── bin/evebox.rs              # 程序入口 + CLI 参数定义
│   ├── cli/                       # CLI 子命令实现
│   │   ├── agent.rs               # Agent 模式
│   │   ├── config/users.rs        # 用户管理 CLI
│   │   ├── elastic/               # ES 管理 CLI
│   │   └── sqlite/                # SQLite 管理 CLI
│   ├── server/                    # HTTP 服务器
│   │   ├── main.rs                # Axum 路由、Session、中间件
│   │   ├── session.rs             # Session 管理
│   │   └── api/                   # REST API 处理
│   │       ├── mod.rs             # 路由注册
│   │       ├── admin.rs           # 管理员 API (用户CRUD)
│   │       ├── login.rs           # 登录/登出
│   │       ├── alerts.rs          # 告警查询
│   │       └── ...
│   ├── sqlite/configdb.rs         # 配置数据库操作（用户、过滤器）
│   └── elastic/                   # ES 客户端
├── webapp/                        # 前端 SolidJS 源码
│   ├── src/
│   │   ├── App.tsx                # 路由 + 认证守卫
│   │   ├── Login.tsx              # 登录页面
│   │   ├── pages/admin/
│   │   │   ├── Admin.tsx          # Admin 首页
│   │   │   ├── AdminUsers.tsx     # 用户管理页面
│   │   │   ├── AdminFilters.tsx   # 过滤器管理
│   │   │   └── AdminElastic.tsx   # ES 管理
│   │   └── ...
│   └── vite.config.ts             # Vite 配置（含 API 代理）
├── resources/
│   ├── webapp/                    # 编译后的前端文件
│   └── configdb/migrations/       # 数据库迁移 SQL
├── examples/                      # 示例配置文件
│   ├── evebox.yaml                # 服务器配置示例
│   └── agent.yaml                 # Agent 配置示例
├── tools/
│   └── suricata_simulator.py     # Suricata EVE 事件模拟器
├── data/                          # 运行数据（gitignore）
│   ├── suricata-eve.json          # 模拟器输出的 EVE 文件
│   └── config.sqlite              # 配置数据库
├── docker/                        # Docker 构建文件
├── control.ps1                    # PowerShell 统一控制脚本
├── control.sh                     # Git Bash 控制脚本
├── start-all.ps1                  # 独立启动脚本
├── stop-all.ps1                   # 独立停止脚本
├── start.txt                      # 开发环境启动说明
├── Makefile                       # 构建脚本
├── DEPLOY.md                      # 本文档 Markdown 版本
└── Cargo.toml                     # Rust 项目配置''')

doc.add_page_break()

# ==================== 11. 常见问题 ====================
add_heading_styled('11. 常见问题', 1)

add_heading_styled('Q1: 启动后 "Failed to fetch" 错误？', 2)
add_para('原因：前端无法连接到后端 API。')
add_bullet('检查后端是否运行：curl http://127.0.0.1:5636/api/version')
add_bullet('如果使用 npm run dev（端口 3636），确保后端也在 5636 端口运行')
add_bullet('如果使用生产模式（端口 5636），确保是直接访问后端')
add_bullet('按 F12 打开浏览器控制台查看详细错误')

add_heading_styled('Q2: \"admin role required\" 错误？', 2)
add_para('表示当前登录用户没有管理员权限。解决方法：')
add_bullet('使用 admin 账号登录')
add_bullet('如果忘记 admin 密码，用 CLI 重置：')
add_code('./target/debug/evebox --data-directory ./data config users passwd admin')

add_heading_styled('Q3: 点击 Create 按钮没有反应？', 2)
add_para('已在最新版本中修复。原因是错误提示被 Modal 弹窗遮挡，现在错误信息直接显示在弹窗内部。如果仍然出现：')
add_bullet('确保已拉取最新代码并重新编译（cargo build）')
add_bullet('强制刷新浏览器（Ctrl+Shift+R）清除旧的 JS 缓存')

add_heading_styled('Q4: 忘记管理员密码？', 2)
add_para('通过 CLI 重置（需指定正确的数据目录）：')
add_code('./target/debug/evebox --data-directory ./data config users passwd admin')
add_para('输入新密码（至少 4 个字符）后即可用新密码登录。')

add_heading_styled('Q5: 如何开启/关闭用户认证？', 2)
add_para('启动时通过以下方式控制：')
add_bullet('命令行：加 --no-auth 禁用认证，不加则默认启用')
add_bullet('配置文件：设置 authentication.required: true/false')
add_bullet('环境变量：EVEBOX_AUTHENTICATION_REQUIRED=true')

add_heading_styled('Q6: 端口 5636 被占用？', 2)
add_para('control.ps1 脚本会自动查找 5636-5649 范围内的空闲端口。手动启动时可通过 -p 参数指定其他端口：')
add_code('./target/debug/evebox server -p 15637 ...')

add_heading_styled('Q7: Elasticsearch 启动失败？', 2)
add_bullet('检查 Java JDK 17 是否安装：java -version')
add_bullet('确认 JDK 路径正确：C:\\Program Files\\Java\\jdk-17.0.18')
add_bullet('检查 elasticsearch/elasticsearch-7.17.28 目录是否存在')
add_bullet('查看 ES 日志：elasticsearch/elasticsearch-7.17.28/logs/')

add_heading_styled('Q8: 构建失败？', 2)
add_bullet('Rust 版本 >= 1.85.0：rustup update')
add_bullet('Node.js 版本 >= v18：node --version')
add_bullet('前端构建前先创建 gitrev.ts：cd webapp && echo "export const GIT_REV = \"$(git rev-parse --short HEAD)\";" > src/gitrev.ts')
add_bullet('清理重新构建：cargo clean && make')
add_bullet('Windows 上如缺少 CMake，安装 Visual Studio Build Tools 或 CMake')

add_heading_styled('Q9: 如何查看 config.sqlite 数据库内容？', 2)
add_code('''# 使用 sqlite3 命令行工具
sqlite3 data/config.sqlite

# 查看用户
SELECT uuid, username, role FROM users WHERE username != "__system__";

# 查看会话
SELECT * FROM sessions;

# 查看过滤器
SELECT * FROM filters;''')

add_heading_styled('Q10: 数据目录（data/）中有哪些文件？', 2)
add_table(
    ['文件', '说明'],
    [
        ['config.sqlite', '配置数据库（用户、过滤器、会话、JA4指纹）'],
        ['events.sqlite', '事件数据库（仅 SQLite 模式）'],
        ['suricata-eve.json', 'Suricata 模拟器生成的 EVE 事件'],
        ['*.bookmark', 'EVE 文件读取位置标记'],
        ['*.pem', 'TLS 自签名证书（自动生成）'],
    ]
)

# ==================== Save ====================
from datetime import datetime
desktop = os.path.join(os.path.expanduser('~'), 'Desktop')
ts = datetime.now().strftime('%Y%m%d_%H%M%S')
filepath = os.path.join(desktop, f'EveBox文档_{ts}.docx')
doc.save(filepath)
print(f'Document saved to: {filepath}')
